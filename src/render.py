"""
MBR Automation — Fase 4: Render
Builds the MBR .docx from the data pack (fetch.py) + the analysis JSON,
following DOC_CONSTRUCTION.md (section order, 2x2 chart grids, table templates).

Usage:
    python src/render.py
"""
import os
import sys
import io
import json
import math

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import config as C

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import matplotlib.font_manager as fm
import numpy as np

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# Design tokens (Nuvemshop Insti Styleguide — DOC_CONSTRUCTION §0.3/0.4)
# ---------------------------------------------------------------------------
C_BRAND      = "0050C3"
C_BRAND_DARK = "171E43"
C_ALT        = "F5FAFF"
C_TOTAL      = "F3F3F3"
C_BORDER     = "E7E8E9"
C_SEP        = "C0C5CC"
C_TEXT       = "343537"
C_WHITE      = "FFFFFF"
C_GRAY       = "A0A3A6"
C_GREEN      = "0CA76B"
C_AMBER      = "9A7000"
C_RED        = "CC3333"

FONT_TEXT  = "Arial"
FONT_TABLE = "Roboto"

CURRENCY = {"BR": "R$", "AR": "$", "MX": "$"}.get(C.COUNTRY, "$")

# Fonts for matplotlib
for _ttf in ("Roboto-Regular.ttf", "Roboto-Bold.ttf"):
    _p = os.path.join(C.FONTS_DIR, _ttf)
    if os.path.exists(_p):
        fm.fontManager.addfont(_p)
_CHART_FONT = "Roboto" if any("Roboto" in f.name for f in fm.fontManager.ttflist) else "DejaVu Sans"
_CS = 8  # chart base font size
_HALF_W, _CH_H, _DPI = 8.1, 7.5, 150


# ---------------------------------------------------------------------------
# Value / format helpers
# ---------------------------------------------------------------------------
def fv(x, d=0.0):
    """Float value from possibly-string/None JSON field."""
    if x is None: return d
    try:
        f = float(x)
        return d if math.isnan(f) else f
    except (ValueError, TypeError):
        return d

def fmt_k(v):
    v = fv(v)
    if abs(v) >= 1_000_000: return f"{v/1_000_000:.1f}M"
    if abs(v) >= 1_000:     return f"{v/1_000:.1f}K"
    return f"{v:,.0f}"

def fmt_int(v):
    return f"{fv(v):,.0f}"

def fmt_money(v):
    v = fv(v)
    if abs(v) >= 1_000_000: return f"{CURRENCY} {v/1_000_000:.1f}M"
    if abs(v) >= 1_000:     return f"{CURRENCY} {v/1_000:.0f}K"
    return f"{CURRENCY} {v:,.0f}"

def fmt_pct(v, dec=1):
    if v is None: return "—"
    return f"{fv(v)*100:.{dec}f}%"

def fmt_signed_pct(v, dec=1):
    if v is None: return "—"
    return f"{fv(v)*100:+.{dec}f}%"

def fmt_pp(v, dec=1):
    """Percentage-point delta (input is a fraction delta, e.g. -0.016 -> -1.6pp)."""
    if v is None: return "—"
    return f"{fv(v)*100:+.{dec}f}pp"

def mom_color(v, inverted=False):
    if v is None: return C_TEXT
    v = fv(v)
    good = (v >= 0) if not inverted else (v <= 0)
    bad  = (v < -0.10) if not inverted else (v > 0.10)
    if good: return C_GREEN
    if bad:  return C_RED
    return C_AMBER

def month_full(label):
    from datetime import datetime
    try: return datetime.strptime(label + "-01", "%Y-%m-%d").strftime("%b %Y")
    except Exception: return label


# ---------------------------------------------------------------------------
# docx helpers
# ---------------------------------------------------------------------------
def init_doc():
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(1.5)
        s.top_margin = s.bottom_margin = Cm(1.5)
    normal = doc.styles["Normal"]
    normal.font.name = FONT_TEXT
    normal.font.size = Pt(10)
    return doc

def doc_title(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.name = FONT_TEXT; r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string(C_BRAND_DARK)

def section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True
    r.font.name = FONT_TEXT; r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(C_BRAND)

def subsection(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True
    r.font.name = FONT_TEXT; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(C_BRAND_DARK)

def body(doc, text, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = FONT_TEXT; r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(C_TEXT)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = FONT_TEXT; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(C_TEXT)

def label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text); r.bold = True; r.italic = True
    r.font.name = FONT_TEXT; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(C_BRAND_DARK)

def spacer(doc, pt=6):
    doc.add_paragraph().paragraph_format.space_after = Pt(pt)


# ---- table cell styling ----
def _tcPr(cell):
    tc = cell._tc
    p = tc.find(qn("w:tcPr"))
    if p is None:
        p = OxmlElement("w:tcPr"); tc.insert(0, p)
    return p

def shade(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    _tcPr(cell).append(shd)

def no_borders(table):
    for row in table.rows:
        for cell in row.cells:
            b = OxmlElement("w:tcBorders")
            for s in ("top", "left", "bottom", "right", "insideH", "insideV"):
                e = OxmlElement(f"w:{s}"); e.set(qn("w:val"), "none"); b.append(e)
            _tcPr(cell).append(b)

def write_cell(cell, text, bold=False, color=None, size=9, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.paragraphs[0].clear() if hasattr(cell.paragraphs[0], "clear") else None
    p = cell.paragraphs[0]
    for r in list(p.runs): r.text = ""
    p.alignment = align
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(str(text)); r.bold = bold
    r.font.name = FONT_TABLE; r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color or C_TEXT)

def table_header(table, headers):
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        shade(c, C_BRAND)
        write_cell(c, h, bold=True, color=C_WHITE, size=9)
    # repeat header row on page breaks
    trPr = table.rows[0]._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true"); trPr.append(th)
    # prevent any single row from splitting across pages
    for row in table.rows:
        rp = row._tr.get_or_add_trPr()
        rp.append(OxmlElement("w:cantSplit"))


# ---------------------------------------------------------------------------
# Chart helpers
# ---------------------------------------------------------------------------
def _style(ax):
    for s in ("top", "right", "left"): ax.spines[s].set_visible(False)
    ax.spines["bottom"].set_color("#" + C_BORDER)
    ax.tick_params(axis="both", labelsize=_CS, colors="#" + C_TEXT)
    ax.tick_params(axis="y", length=0)
    ax.yaxis.grid(True, color="#" + C_BORDER, linewidth=0.5, zorder=0)
    ax.set_axisbelow(True)

def _xlabels(months):
    from datetime import datetime
    out, prev = [], None
    for m in months:
        try:
            dt = datetime.strptime(m + "-01", "%Y-%m-%d")
            out.append(dt.strftime("%b '%y") if dt.year != prev else dt.strftime("%b"))
            prev = dt.year
        except Exception:
            out.append(m)
    return out

def _xaxis(ax, months):
    """Sparse x labels: every 3rd month, always including the last."""
    n = len(months)
    idx = sorted(set(list(range(n - 1, -1, -3)) + [n - 1])) if n else []
    labels = _xlabels(months)
    ax.set_xticks(idx)
    ax.set_xticklabels([labels[i] for i in idx], fontsize=_CS - 1, rotation=0)

def _buf(fig):
    b = io.BytesIO()
    fig.savefig(b, format="png", dpi=_DPI, bbox_inches="tight", facecolor="white")
    b.seek(0); plt.close(fig)
    return b

def chart_bars(months, vals, title, ylabel="", color="#2E7D32", pct=False, money=False):
    plt.rcParams["font.family"] = _CHART_FONT
    fig, ax = plt.subplots(figsize=(_HALF_W/2.54, _CH_H/2.54)); _style(ax)
    x = np.arange(len(months))
    colors = [color]*len(vals)
    if len(vals): colors[-1] = "#1B5E20" if color == "#2E7D32" else color
    bars = ax.bar(x, vals, 0.7, color=colors, zorder=3)
    if len(vals):
        v = vals[-1]
        txt = (f"{v:.1f}%" if pct else (fmt_money(v) if money else fmt_k(v)))
        ax.text(x[-1], v, txt, ha="center", va="bottom", fontsize=_CS-1, color="#"+C_TEXT)
    _xaxis(ax, months)
    if pct: ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_: f"{v:.0f}%"))
    else:   ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_: fmt_k(v)))
    ax.set_title(title, fontsize=_CS, color="#"+C_TEXT, pad=4)
    fig.tight_layout(); return _buf(fig)

def chart_line(months, vals, title, plan=None, pct=False, money=False, fill=False):
    plt.rcParams["font.family"] = _CHART_FONT
    fig, ax = plt.subplots(figsize=(_HALF_W/2.54, _CH_H/2.54)); _style(ax)
    x = np.arange(len(months))
    if fill: ax.fill_between(x, vals, alpha=0.12, color="#"+C_BRAND, zorder=2)
    ax.plot(x, vals, "o-", color="#"+C_BRAND, linewidth=1.5, markersize=2.5, zorder=4)
    if plan is not None and any(plan):
        ax.plot(x, plan, "--", color="#"+C_SEP, linewidth=1.0, zorder=3, label="Plan")
        ax.legend(loc="best", fontsize=_CS-1, frameon=False)
    if len(vals):
        v = vals[-1]
        txt = (f"{v:.1f}%" if pct else (fmt_money(v) if money else fmt_k(v)))
        ax.annotate(txt, xy=(x[-1], v), xytext=(0,6), textcoords="offset points",
                    ha="center", fontsize=_CS-1, color="#"+C_BRAND)
    _xaxis(ax, months)
    if pct: ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_: f"{v:.0f}%"))
    else:   ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_: fmt_k(v)))
    ax.set_title(title, fontsize=_CS, color="#"+C_TEXT, pad=4)
    fig.tight_layout(); return _buf(fig)

def chart_netadds(months, vals, title):
    plt.rcParams["font.family"] = _CHART_FONT
    fig, ax = plt.subplots(figsize=(_HALF_W/2.54, _CH_H/2.54)); _style(ax)
    x = np.arange(len(months))
    colors = ["#0CA76B" if v >= 0 else "#CC3333" for v in vals]
    ax.bar(x, vals, 0.7, color=colors, zorder=3)
    ax.axhline(0, color="#"+C_BORDER, linewidth=0.8)
    _xaxis(ax, months)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_: fmt_k(v)))
    ax.set_title(title, fontsize=_CS, color="#"+C_TEXT, pad=4)
    fig.tight_layout(); return _buf(fig)

def chart_stacked(months, series, title, colors):
    plt.rcParams["font.family"] = _CHART_FONT
    fig, ax = plt.subplots(figsize=(_HALF_W/2.54, _CH_H/2.54)); _style(ax)
    x = np.arange(len(months)); bottom = np.zeros(len(months))
    for (name, vals), col in zip(series.items(), colors):
        ax.bar(x, vals, 0.7, bottom=bottom, color=col, label=name, zorder=3)
        bottom += np.array(vals)
    _xaxis(ax, months)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_: fmt_k(v)))
    ax.set_title(title, fontsize=_CS, color="#"+C_TEXT, pad=4)
    ax.legend(loc="upper left", fontsize=_CS-2, frameon=False)
    fig.tight_layout(); return _buf(fig)

def chart_combo(months, bars, lines, title, bar_label, line_labels, line_colors):
    plt.rcParams["font.family"] = _CHART_FONT
    fig, ax = plt.subplots(figsize=(_HALF_W/2.54, _CH_H/2.54)); _style(ax)
    x = np.arange(len(months))
    ax.bar(x, bars, 0.7, color="#B0BEC5", label=bar_label, zorder=2)
    for vals, lab, col in zip(lines, line_labels, line_colors):
        ax.plot(x, vals, "o-", color=col, linewidth=1.2, markersize=2, label=lab, zorder=4)
    _xaxis(ax, months)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v,_: fmt_k(v)))
    ax.set_title(title, fontsize=_CS, color="#"+C_TEXT, pad=4)
    ax.legend(loc="upper right", fontsize=_CS-2, frameon=False)
    fig.tight_layout(); return _buf(fig)

def add_chart_grid(doc, bufs):
    """bufs: list of up to 4 BytesIO. Lays out 2x2 in an invisible table."""
    rows = (len(bufs) + 1) // 2
    table = doc.add_table(rows=rows, cols=2)
    no_borders(table)
    half = Cm(_HALF_W)
    for i, buf in enumerate(bufs):
        cell = table.cell(i // 2, i % 2); cell.width = half
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(buf, width=half)
    spacer(doc, 4)


# ---------------------------------------------------------------------------
# Generic analysis block
# ---------------------------------------------------------------------------
def analysis_block(doc, a):
    if a.get("results"):
        label(doc, "Results")
        body(doc, a["results"], size=9)
    if a.get("gap"):
        label(doc, "Gap & Hypothesis")
        for b in a["gap"]: bullet(doc, b)
    if a.get("moving_forward"):
        label(doc, "Moving Forward")
        for b in a["moving_forward"]: bullet(doc, b)
    spacer(doc)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------
def build_tldr(doc, an):
    t = an["tldr"]
    section_title(doc, "1. TL;DR")
    subsection(doc, "Executive Summary"); body(doc, t["executive_summary"])
    subsection(doc, "The Funnel Narrative"); body(doc, t["funnel_narrative"])
    subsection(doc, "Base Health & Retention"); body(doc, t["base_health"])
    subsection(doc, "Business & Financial Impact"); body(doc, t["financial_impact"])
    subsection(doc, "Critical Discussion Topics")
    for x in t["critical_topics"]: bullet(doc, x)
    subsection(doc, "Prioritized Action Plan")
    label(doc, "Immediate Fixes (30d)")
    for x in t["action_plan"]["immediate"]: bullet(doc, x)
    label(doc, "Strategic Bets (Q+1)")
    for x in t["action_plan"]["strategic"]: bullet(doc, x)
    spacer(doc)


T1_HEADERS = ["KPI", "Real", "Plan", "%Achv", "YoY%", "YTD", "%Achv YTD", "YoY% YTD"]

def write_kpi_row(table, ri, name, kpi, kind="num", fmt=fmt_k, inverted=False):
    """T1 row (DOC_CONSTRUCTION): KPI|Real|Plan|%Achv|YoY%|YTD|%Achv YTD|YoY% YTD.
    kind: num | money | ratio | blocked."""
    cells = table.rows[ri].cells
    bg = C_ALT if ri % 2 == 0 else C_WHITE
    for c in cells: shade(c, bg)
    write_cell(cells[0], name, color=C_TEXT, align=WD_ALIGN_PARAGRAPH.LEFT)
    if kind == "blocked":
        for ci in range(1, 8): write_cell(cells[ci], "—", color=C_GRAY)
        return
    if kind == "ratio":
        write_cell(cells[1], fmt_pct(kpi.get("real")))
        write_cell(cells[2], fmt_pct(kpi["plan"]) if kpi.get("plan") else "—")
        write_cell(cells[3], "—")
        write_cell(cells[4], "—")
        write_cell(cells[5], fmt_pct(kpi.get("ytd")) if kpi.get("ytd") is not None else "—")
        write_cell(cells[6], "—")
        write_cell(cells[7], "—")
        return
    # num / money
    att = kpi.get("attainment"); att_ytd = kpi.get("attainment_ytd")
    yoy = kpi.get("yoy_pct"); yoy_ytd = kpi.get("yoy_ytd_pct")
    write_cell(cells[1], fmt(kpi.get("real")))
    write_cell(cells[2], fmt(kpi["plan"]) if kpi.get("plan") else "—")
    write_cell(cells[3], fmt_pct(att) if att is not None else "—",
               color=mom_color((att - 1) if att is not None else None, inverted))
    write_cell(cells[4], fmt_signed_pct(yoy) if yoy is not None else "—", color=mom_color(yoy, inverted))
    write_cell(cells[5], fmt(kpi["ytd"]) if kpi.get("ytd") is not None else "—")
    write_cell(cells[6], fmt_pct(att_ytd) if att_ytd is not None else "—",
               color=mom_color((att_ytd - 1) if att_ytd is not None else None, inverted))
    write_cell(cells[7], fmt_signed_pct(yoy_ytd) if yoy_ytd is not None else "—", color=mom_color(yoy_ytd, inverted))


def build_main_kpis(doc, dp, an):
    section_title(doc, "2. Main KPIs")
    mac = dp["acquisition"]["macro"]; fin = dp["financial"]["gmv"]; cm = dp["company_metrics"]["base"]
    br = dp["brand"]["macro"]
    # (name, kpi, kind, fmt, inverted) — DOC_CONSTRUCTION §2.1, top-of-funnel -> financial
    rows = [
        ("Branded Searches",   br["branded_searches"],     "num",   fmt_k,     False),
        ("Sessions",           mac["sessions"],            "num",   fmt_k,     False),
        ("Trials",             mac["trials"],              "num",   fmt_k,     False),
        ("CVR Session->Trial", mac["cvr_session_trial"],   "ratio", None,      False),
        ("New Payments",       mac["new_payments"],        "num",   fmt_k,     False),
        ("CVR Trial->NP",      mac["cvr_trial_np"],        "ratio", None,      False),
        ("New Sellers",        mac["new_sellers"],         "num",   fmt_k,     False),
        ("Net Adds",           cm["net_adds"],             "num",   fmt_int,   False),
        ("Merchant Base",      cm["merchant_base"],        "num",   fmt_k,     False),
        ("GMV (local)",        fin["gmv_lc"],              "money", fmt_money, False),
        ("Orders",             fin["orders"],              "num",   fmt_k,     False),
        ("% Seller",           {},                         "blocked", None,    False),
        ("GMV per Seller",     {},                         "blocked", None,    False),
    ]
    table = doc.add_table(rows=len(rows)+1, cols=8); table.style = "Table Grid"
    table_header(table, T1_HEADERS)
    for i, (name, kpi, kind, fmt, inv) in enumerate(rows, 1):
        write_kpi_row(table, i, name, kpi, kind, fmt or fmt_k, inv)
    spacer(doc)
    # 2x2 charts: Trials, NPs, Net Adds, GMV (local)
    fh = dp["acquisition"]["history"]["funnel"]; gh = dp["financial"]["history"]["gmv"]; ch = dp["company_metrics"]["history"]
    m = [r["month_label"] for r in fh]
    add_chart_grid(doc, [
        chart_bars(m, [fv(r["trials"]) for r in fh], "Trials"),
        chart_bars(m, [fv(r["nps"]) for r in fh], "New Payments", color="#2E7D32"),
        chart_netadds([r["month_label"] for r in ch], [fv(r["net_adds"]) for r in ch], "Net Adds"),
        chart_line([r["month_label"] for r in gh], [fv(r["gmv_lc"]) for r in gh], "GMV (local)", money=True, fill=True),
    ])


def build_brand(doc, dp, an):
    section_title(doc, "3. Brand & Comms")
    b = dp["brand"]; mac = b["macro"]
    # 3.1 charts
    subsection(doc, "3.1 Macro Results")
    sh = b["history"]["search"]; m = [r["month_label"] for r in sh]
    add_chart_grid(doc, [
        chart_bars(m, [fv(r["branded_searches"]) for r in sh], "Branded Searches"),
        chart_bars(m, [fv(r["clicks"])/fv(r["branded_searches"],1)*100 if fv(r["branded_searches"]) else 0 for r in sh], "Branded CTR %", color="#81C784", pct=True),
        chart_combo(m, [fv(r["total_market"]) for r in sh],
                    [[fv(r["branded_kwp"]) for r in sh], [fv(r["nonbranded_kwp"]) for r in sh]],
                    "Total Market", "Total", ["Branded", "Non-Branded"], ["#"+C_BRAND, "#CC3333"]),
        chart_line(m, [fv(r["branded_kwp"])/fv(r["d2c_kwp"],1)*100 if fv(r["d2c_kwp"]) else 0 for r in sh], "Share of Market", pct=True),
    ])
    # 3.1 table (T1)
    table = doc.add_table(rows=5, cols=8); table.style = "Table Grid"
    table_header(table, T1_HEADERS)
    write_kpi_row(table, 1, "Branded Searches", mac["branded_searches"], "num", fmt_k)
    write_kpi_row(table, 2, "Branded CTR %", mac["branded_ctr"], "ratio")
    write_kpi_row(table, 3, "Share of Market", mac["share_of_market"], "ratio")
    write_kpi_row(table, 4, "Total Market", mac["total_market"], "num", fmt_k)
    spacer(doc)
    # 3.2 PR
    subsection(doc, "3.2 PR")
    pr_m = b["pr"]["month"]; pr_p = b["pr"]["prior_month"]
    pr_rows = [("Gold", "gold"), ("Tier 1", "tier1"), ("Tier 2", "tier2"), ("Total", "total")]
    table = doc.add_table(rows=len(pr_rows)+1, cols=4); table.style = "Table Grid"
    table_header(table, ["Tier", "Real", "M-1", "MoM%"])
    for ri, (nm, key) in enumerate(pr_rows, 1):
        cells = table.rows[ri].cells; is_total = (nm == "Total")
        bg = C_TOTAL if is_total else (C_ALT if ri % 2 == 0 else C_WHITE)
        for c in cells: shade(c, bg)
        real, prior = fv(pr_m.get(key)), fv(pr_p.get(key))
        mom = (real-prior)/abs(prior) if prior else None
        write_cell(cells[0], nm, bold=is_total, color=C_TEXT, align=WD_ALIGN_PARAGRAPH.LEFT)
        write_cell(cells[1], fmt_int(real), bold=is_total)
        write_cell(cells[2], fmt_int(prior), bold=is_total)
        write_cell(cells[3], fmt_signed_pct(mom) if mom is not None else "—", color=mom_color(mom))
    body(doc, f"Proactive: {fmt_int(pr_m.get('proactive'))} · Total Reach: "
              f"{'N/A' if fv(pr_m.get('reach'))==0 else fmt_k(pr_m.get('reach'))}", size=9)
    spacer(doc)
    # 3.3 Social
    subsection(doc, "3.3 Social")
    sm = b["social"]["month"]; sp = b["social"]["prior_month"]
    fm_ = b["social"]["followers_month"]; fp = b["social"]["followers_prior_month"]
    soc_rows = [("Total Reach", fv(sm.get("reach_total")), fv(sp.get("reach_total"))),
                ("Organic Reach", fv(sm.get("reach_organic")), fv(sp.get("reach_organic"))),
                ("Paid Reach", fv(sm.get("reach_paid")), fv(sp.get("reach_paid"))),
                ("Collab Reach", fv(sm.get("reach_collab")), fv(sp.get("reach_collab"))),
                ("Interactions", fv(sm.get("interactions")), fv(sp.get("interactions"))),
                ("Net Followers", fv(fm_.get("net_followers")), fv(fp.get("net_followers"))),
                ("Profile Followers", fv(fm_.get("followers")), fv(fp.get("followers")))]
    table = doc.add_table(rows=len(soc_rows)+1, cols=4); table.style = "Table Grid"
    table_header(table, ["Metric", "Real", "M-1", "MoM%"])
    for ri, (nm, real, prior) in enumerate(soc_rows, 1):
        cells = table.rows[ri].cells; bg = C_ALT if ri % 2 == 0 else C_WHITE
        for c in cells: shade(c, bg)
        mom = (real-prior)/abs(prior) if prior else None
        write_cell(cells[0], nm, color=C_TEXT, align=WD_ALIGN_PARAGRAPH.LEFT)
        write_cell(cells[1], fmt_k(real)); write_cell(cells[2], fmt_k(prior))
        write_cell(cells[3], fmt_signed_pct(mom) if mom is not None else "—", color=mom_color(mom))
    spacer(doc)
    analysis_block(doc, an["brand"])


def build_acquisition(doc, dp, an):
    section_title(doc, "4. Acquisition")
    mac = dp["acquisition"]["macro"]
    subsection(doc, "4.1 Macro Results")
    fh = dp["acquisition"]["history"]["funnel"]; sh = dp["acquisition"]["history"]["sessions"]
    m = [r["month_label"] for r in fh]
    cvr_hist = [fv(r["nps"])/fv(r["trials"],1)*100 if fv(r["trials"]) else 0 for r in fh]
    add_chart_grid(doc, [
        chart_bars([r["month_label"] for r in sh], [fv(r["sessions"]) for r in sh], "Sessions", color="#5A6BC0"),
        chart_bars(m, [fv(r["trials"]) for r in fh], "Trials"),
        chart_bars(m, [fv(r["nps"]) for r in fh], "New Payments", color="#2E7D32"),
        chart_line(m, cvr_hist, "CVR Trial->NP", pct=True),
    ])
    # 4.1 macro table T1
    rows = [("Sessions", mac["sessions"], fmt_k, False),
            ("Trials", mac["trials"], fmt_k, False),
            ("New Payments", mac["new_payments"], fmt_k, False),
            ("New Sellers", mac["new_sellers"], fmt_k, False),
            ("QLs", mac["qls"], fmt_k, False)]
    table = doc.add_table(rows=len(rows)+1, cols=8); table.style = "Table Grid"
    table_header(table, T1_HEADERS)
    for i, (nm, kpi, fmt, inv) in enumerate(rows, 1):
        write_kpi_row(table, i, nm, kpi, "num", fmt, inv)
    spacer(doc)
    # 4.2 by source (T3) — grouped L4 / L2
    subsection(doc, "4.2 By Source (Branded / Non-Branded)")
    agg = {}
    for r in dp["acquisition"]["by_source"]:
        src = r["mkt_source"]; l2 = C.level2(src); l4 = C.level4(src)
        d = agg.setdefault((l4, l2), {"trials": 0, "nps": 0, "nss": 0})
        d["trials"] += fv(r["trials"]); d["nps"] += fv(r["nps"]); d["nss"] += fv(r["nss"])
    # order
    order_nb = ["Performance No Brand", "Affiliates", "Partners", "Organic Growth", "Nurturing", "Misc Sources"]
    order_b = ["Performance Brand", "Organic", "Direct"]
    def grp_rows(l4, order):
        items = [(l2, agg[(l4,l2)]) for l2 in order if (l4,l2) in agg]
        # any not in order
        for (gl4,l2), d in agg.items():
            if gl4==l4 and l2 not in order: items.append((l2,d))
        return items
    b_items = grp_rows("Branded", order_b); nb_items = grp_rows("Non-Branded", order_nb)
    total_rows = 1 + (1+len(b_items)) + (1+len(nb_items)) + 1
    table = doc.add_table(rows=total_rows, cols=5); table.style = "Table Grid"
    table_header(table, ["Source", "Trials", "NPs", "NSs", "CVR T->NP"])
    ri = 1
    def subtotal(name, items):
        nonlocal ri
        tt = sum(d["trials"] for _,d in items); tn = sum(d["nps"] for _,d in items); ts = sum(d["nss"] for _,d in items)
        cells = table.rows[ri].cells
        for c in cells: shade(c, C_TOTAL)
        write_cell(cells[0], name, bold=True, color=C_BRAND_DARK, align=WD_ALIGN_PARAGRAPH.LEFT)
        write_cell(cells[1], fmt_k(tt), bold=True); write_cell(cells[2], fmt_k(tn), bold=True)
        write_cell(cells[3], fmt_k(ts), bold=True)
        write_cell(cells[4], fmt_pct(tn/tt) if tt else "—", bold=True)
        ri += 1
        for l2, d in items:
            cells = table.rows[ri].cells; bg = C_ALT if ri % 2 == 0 else C_WHITE
            for c in cells: shade(c, bg)
            write_cell(cells[0], "  " + l2, color=C_TEXT, align=WD_ALIGN_PARAGRAPH.LEFT)
            write_cell(cells[1], fmt_k(d["trials"])); write_cell(cells[2], fmt_k(d["nps"]))
            write_cell(cells[3], fmt_k(d["nss"]))
            write_cell(cells[4], fmt_pct(d["nps"]/d["trials"]) if d["trials"] else "—")
            ri += 1
    subtotal("Branded", b_items); subtotal("Non-Branded", nb_items)
    # grand total
    allt = sum(d["trials"] for d in agg.values()); alln = sum(d["nps"] for d in agg.values()); alls = sum(d["nss"] for d in agg.values())
    cells = table.rows[ri].cells
    for c in cells: shade(c, C_BRAND);
    write_cell(cells[0], "Total", bold=True, color=C_WHITE, align=WD_ALIGN_PARAGRAPH.LEFT)
    write_cell(cells[1], fmt_k(allt), bold=True, color=C_WHITE); write_cell(cells[2], fmt_k(alln), bold=True, color=C_WHITE)
    write_cell(cells[3], fmt_k(alls), bold=True, color=C_WHITE)
    write_cell(cells[4], fmt_pct(alln/allt) if allt else "—", bold=True, color=C_WHITE)
    spacer(doc)
    analysis_block(doc, an["acquisition"])


def build_company_metrics(doc, dp, an):
    section_title(doc, "5. Company Metrics")
    cm = dp["company_metrics"]; base = cm["base"]; comp = cm["components"]; ch = cm["history"]
    subsection(doc, "5.1 Merchant Base & Net Adds")
    m = [r["month_label"] for r in ch]
    add_chart_grid(doc, [
        chart_line(m, [fv(r["merchant_base"]) for r in ch], "Merchant Base", fill=True),
        chart_netadds(m, [fv(r["net_adds"]) for r in ch], "Net Adds"),
        chart_stacked(m, {"First Payments": [fv(r["first_payments"]) for r in ch],
                          "New Phoenix": [fv(r["new_phoenix"]) for r in ch],
                          "Old Phoenix": [fv(r["old_phoenix"]) for r in ch]},
                      "Inflow Components", ["#2E7D32", "#66BB6A", "#A5D6A7"]),
        chart_stacked(m, {"Churns": [fv(r["churns"]) for r in ch],
                          "Downsell": [fv(r["downsell_freemium"]) for r in ch]},
                      "Outflow Components", ["#C62828", "#EF9A9A"]),
    ])
    # summary table T2 + YoY
    rows = [("Net Adds", base["net_adds"], fmt_int, True),
            ("First Payments", base["first_payments"], fmt_k, False),
            ("Churn + Downgrade", base["churn_and_downgrade"], fmt_k, True),
            ("Merchant Base (EOP)", base["merchant_base"], fmt_k, True)]
    table = doc.add_table(rows=len(rows)+1, cols=5); table.style = "Table Grid"
    table_header(table, ["Metric", "Real", "M-1", "MoM%", "YoY%"])
    for i, (nm, kpi, fmt, bold) in enumerate(rows, 1):
        cells = table.rows[i].cells; bg = C_TOTAL if bold else (C_ALT if i % 2 == 0 else C_WHITE)
        for c in cells: shade(c, bg)
        write_cell(cells[0], nm, bold=bold, color=C_TEXT, align=WD_ALIGN_PARAGRAPH.LEFT)
        write_cell(cells[1], fmt(kpi.get("real")), bold=bold)
        write_cell(cells[2], fmt(kpi.get("m1")))
        write_cell(cells[3], fmt_signed_pct(kpi.get("mom")) if kpi.get("mom") is not None else "—", color=mom_color(kpi.get("mom")))
        write_cell(cells[4], fmt_signed_pct(kpi.get("yoy_pct")) if kpi.get("yoy_pct") is not None else "—", color=mom_color(kpi.get("yoy_pct")))
    spacer(doc)
    analysis_block(doc, an["company_metrics"])


def build_financial(doc, dp, an):
    section_title(doc, "6. Financial")
    g = dp["financial"]["gmv"]; gh = dp["financial"]["history"]["gmv"]
    subsection(doc, "6.1 GMV")
    m = [r["month_label"] for r in gh]
    add_chart_grid(doc, [
        chart_line(m, [fv(r["gmv_lc"]) for r in gh], "GMV (local)", money=True, fill=True),
        chart_bars(m, [fv(r["gmv_usd"]) for r in gh], "GMV (USD)", color="#2E7D32", money=True),
    ])
    rows = [("GMV (local)", g["gmv_lc"], fmt_money),
            ("GMV (USD)", g["gmv_usd"], fmt_money),
            ("Orders", g["orders"], fmt_k),
            ("Avg Ticket", {"real": g["avg_ticket"]["real"], "m1": g["avg_ticket"]["m1"]}, fmt_money)]
    table = doc.add_table(rows=len(rows)+1, cols=5); table.style = "Table Grid"
    table_header(table, ["KPI", "Real", "M-1", "MoM%", "YoY%"])
    for i, (nm, kpi, fmt) in enumerate(rows, 1):
        cells = table.rows[i].cells; bg = C_ALT if i % 2 == 0 else C_WHITE
        for c in cells: shade(c, bg)
        write_cell(cells[0], nm, color=C_TEXT, align=WD_ALIGN_PARAGRAPH.LEFT)
        write_cell(cells[1], fmt(kpi.get("real")))
        write_cell(cells[2], fmt(kpi.get("m1")) if kpi.get("m1") is not None else "—")
        write_cell(cells[3], fmt_signed_pct(kpi.get("mom")) if kpi.get("mom") is not None else "—", color=mom_color(kpi.get("mom")))
        write_cell(cells[4], fmt_signed_pct(kpi.get("yoy_pct")) if kpi.get("yoy_pct") is not None else "—", color=mom_color(kpi.get("yoy_pct")))
    spacer(doc)
    analysis_block(doc, an["financial"])


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    dp = json.load(open(C.DATA_PACK_PATH, encoding="utf-8"))
    an_path = os.path.join(C.OUTPUT_DIR, f"analysis_{C.COUNTRY}_{C.MONTH_LABEL}.json")
    an = json.load(open(an_path, encoding="utf-8"))

    doc = init_doc()
    doc_title(doc, f"MBR · SMB {C.COUNTRY} · {month_full(C.MONTH_LABEL)}")
    spacer(doc)
    build_tldr(doc, an)
    build_main_kpis(doc, dp, an)
    build_brand(doc, dp, an)
    build_acquisition(doc, dp, an)
    build_company_metrics(doc, dp, an)
    build_financial(doc, dp, an)

    doc.save(C.DOCX_PATH)
    print(f"[render] Saved: {C.DOCX_PATH}")


if __name__ == "__main__":
    main()
